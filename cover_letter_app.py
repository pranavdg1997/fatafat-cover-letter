import warnings
warnings.filterwarnings("ignore")
import streamlit as st
import pandas as pd
import numpy as np
import datetime
from docx import Document
import yaml

#format datetime into string
def format_dt(dt):
	return("{}-{}-{} {}:{}".format(dt.month,dt.day,dt.year,dt.hour,dt.minute))  



st.title('Fatafat Cover Letter')
st.write('This app lets you speed up making your custom cover letter. You can customize & replace Company name and Position in your cover letter. You can also keep track of positions applied so far.')



configs = yaml.load(open("configs.yml"))
document = Document(configs["parent_filepath"])
history_df = pd.read_csv(configs["history_file"])

para_dict = {}
st.markdown("## Fast Edit features")
with st.beta_container():
    company_name = st.text_input("Enter company name")
    profile_name = st.text_input("Enter profile name")
    ri = st.slider(
        "Adjust paragraph number (make sure original para & replaced para match except for the above two inputs)",
        0,
        len(document.paragraphs),
        7,
        1)

    replaced_para = configs["replace_text"].format(company_name,profile_name)
    print(replaced_para)

    
    st.markdown("#### Original paragraph:")
    if(document.paragraphs[ri].text==''):
        st.write("<blank para>")
    else:
        st.write(document.paragraphs[ri].text)
    st.markdown("#### Replaced paragraph:")
    replaced_para_text = st.write(replaced_para)

    button1 = st.button(key="button1",
        label="Generate Cover Letter!"
        )
    if(button1):
        document.paragraphs[ri].text = replaced_para
        document.save(configs["output_filepath"])
        new_entry = pd.DataFrame({"Company Name":company_name,
        	"Profile":profile_name,
        	"Timestamp":format_dt(datetime.datetime.now())},index=[0])
        history_df = pd.concat([new_entry, history_df],axis=0)
        st.write("File written to {}".format(configs["output_filepath"]))
        history_df.to_csv(configs["history_file"],index=False)

st.markdown("## Previous applications")
with st.beta_container():
	st.write(history_df)

